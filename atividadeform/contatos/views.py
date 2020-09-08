from django.shortcuts import render, redirect, get_object_or_404
from .models import Produtos, ListaMaterial, Fornecedor, Grupos, Projeto, DocFiles
from .forms import FormularioContato, FormularioLista, FormularioFornecedor, FormularioProjeto, NameForm
import openpyxl
from openpyxl.styles import Font, colors, Alignment, Border, Side, PatternFill
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse, Http404
import os
from django.conf import settings
from django.core.files import File


def listar_contatos(request):
    busca = request.GET.get('pesquisa',None)

    if busca:
        # contatos = Contatos.objects.all()
        produtos = Produtos.objects.filter(nome__icontains = busca) or Produtos.objects.filter(fabricante__icontains = busca)
    else:
        produtos = Produtos.objects.order_by('nome')

    return render(request, 'contatos.html', {'contatos' : produtos})

# def listar_produtos(request):
#     busca = request.GET.get('pesquisa', None)
#
#     produtos = ListaMaterial.objects.all()
#
#     return render(request, 'formulario_lista.html')

def novo_contato(request):
    form = FormularioContato(request.POST or None)

    if form.is_valid():
        form.save()
        return redirect('lista_contatos')

    return  render(request, 'formulario_contato.html', {'form' : form})

def atualizar_contato(request, id):
    contato = get_object_or_404(Produtos, pk=id)
    form = FormularioContato(request.POST or None, instance = contato)

    if form.is_valid():
        form.save()
        return redirect('lista_contatos')

    return render(request,'formulario_contato.html', {'form' : form})


def atualizar_fornecedor(request, id):
    fornecedor = get_object_or_404(Fornecedor, pk=id)
    form = FormularioFornecedor(request.POST or None, instance=fornecedor)

    if form.is_valid():
        form.save()
        return redirect('lista_fornecedor')

    return render(request, 'formulario_fornecedor.html', {'form': form})

def atualizar_prod_lista(request, id):
    produtos = ListaMaterial.objects.order_by('produto__nome')
    produto = get_object_or_404(ListaMaterial, pk=id)
    form = FormularioLista(request.POST or None, instance=produto)

    if form.is_valid():
        form.save()
        return redirect('adicionar_lista')

    return render(request, 'formulario_lista.html', {'form': form},)


def excluir_produto(request,id):
    produto = get_object_or_404(Produtos, id =  id)
    form = FormularioContato(request.POST or None, request.FILES or None, instance = produto)
    if request.method == 'POST':
        produto.delete()
        return redirect('lista_contatos')

    # if request.method is not 'POST':
    #     post_delete=Contatos.objects.filter(id=id)
    #     post_delete.delete()
    #     return redirect('lista_contatos')

    return render(request, 'confirmar_delete_produto.html', {'produto': produto})

def excluir_lista_produto(request,id):
    produto = get_object_or_404(ListaMaterial, id =  id)
    form = FormularioContato(request.POST or None, request.FILES or None, instance = produto)
    if request.method == 'POST':
        produto.delete()
        return redirect('adicionar_lista')

    # if request.method is not 'POST':
    #     post_delete=Contatos.objects.filter(id=id)
    #     post_delete.delete()
    #     return redirect('lista_contatos')

    return render(request, 'confirmar_delete_produto.html', {'produto': produto})

# def excluir_prod_lista(request,id):
#     produto = get_object_or_404(ListaMaterial, id = id)
#     form = FormularioLista(request.POST or None)
#     produtos = ListaMaterial.objects.all()
#     contatos = Produtos.objects.all()
#
#     if request.method is not 'POST':
#         post_delete=ListaMaterial.objects.filter(id=id)
#         post_delete.delete()
#         produtos = ListaMaterial.objects.all()
#         contatos = Produtos.objects.all()
#         return render(request, 'formulario_lista.html', {'form': form, 'produtos': produtos, 'contatos': contatos})
#     return render(request, 'formulario_lista.html', {'form': form, 'produtos': produtos, 'contatos': contatos})

def novo_fornecedor(request):
    form = FormularioFornecedor(request.POST or None)

    if form.is_valid():
        form.save()
        return redirect('listar_fornecedor')
    return render(request,'formulario_fornecedor.html',{'form': form})


def listar_fornecedor(request):
    fornecedores = Fornecedor.objects.order_by('razao_social')

    return render(request, 'lista_fornecedor.html', {'fornecedores' : fornecedores})

def excluir_fornecedor(request,id):
    if request.method is not 'POST':
        forn_delete=Fornecedor.objects.filter(id=id)
        forn_delete.delete()
        return redirect('lista_fornecedor')


def nova_lista(request):
    form = FormularioLista(request.POST or None)

    if form.is_valid():
        form.save()
        produtos = ListaMaterial.objects.order_by('produto__nome')

        return redirect('adicionar_lista')

    produtos = ListaMaterial.objects.order_by('produto__nome')

    # for i in produtos:
    #     print(i.produto.fabricante)

    return render(request, 'formulario_lista.html', {'form': form,'produtos':produtos})

def excluir_prod_lista(request,id):
    if request.method is not 'POST':
        print(id)
        ListaMaterial.objects.filter(id=id).delete()

        return redirect('adicionar_lista')


def novo_projeto(request):
    listas = ListaMaterial.objects.order_by('produto__nome')

    form = FormularioProjeto(request.POST or None)

    if form.is_valid():
        form.save()
        return redirect('adicionar_lista')
    return render(request, 'formulario_projeto.html', {'form': form})


def gerar_xlsx(nome_doc):

    produtos = Produtos.objects.order_by('nome')
    lista = ListaMaterial.objects.order_by('produto__nome')

    wb = openpyxl.Workbook()
    cont = 2
    planilha = wb.active
    planilha.title = 'PREÇOS'

    planilha['A1'] = 'DESCRIÇÃO'
    planilha['B1'] = 'MODELO'
    planilha['C1'] = 'FABRICANTE'
    planilha['D1'] = 'FORNECEDOR'
    planilha['E1'] = 'QUANTIDADE'
    planilha['F1'] = 'VALOR UNITÁRIO'
    planilha['G1'] = 'DATA COTAÇÃO'
    planilha['H1'] = 'TOTAL'


    #FORMATANDO AS CÉLULAS DO ARQUIVO
    ft_cabecalho = Font(name='Arial', size=12, bold=True, color= 'FFFFFF')
    ft_item = Font(name='Arial',size=10)
    ft_item_negrito = Font(name='Arial', size=10,bold=True)
    ft_item_italico = Font(name='Arial', size=10, italic=True, color='505050')

    alinhamento = Alignment(horizontal='center',vertical='center')
    alinhamentoEsquerda = Alignment(horizontal='left', vertical='center')

    fina = Side(border_style='thin', color='000000')
    bordaInferior = Border(bottom=fina)
    bordaSuperior = Border(top=fina)

    preenchimentoAzul = PatternFill('solid',fgColor='6495ED')
    preenchimentoVerde = PatternFill('solid', fgColor='00FF7F')
    preenchimentoCinza = PatternFill('solid', fgColor='DDDDDD')
    preenchimentoAzulClaro = PatternFill('solid', fgColor='dbe5f1')



    #COMEÇANDO A ESCREVER O ARQUIVO XLSX
    planilha['A1'].font = ft_cabecalho
    planilha['B1'].font = ft_cabecalho
    planilha['C1'].font = ft_cabecalho
    planilha['D1'].font = ft_cabecalho
    planilha['E1'].font = ft_cabecalho
    planilha['F1'].font = ft_cabecalho
    planilha['G1'].font = ft_cabecalho
    planilha['H1'].font = ft_cabecalho

    planilha['A1'].border = bordaInferior
    planilha['B1'].border = bordaInferior
    planilha['C1'].border = bordaInferior
    planilha['D1'].border = bordaInferior
    planilha['E1'].border = bordaInferior
    planilha['F1'].border = bordaInferior
    planilha['G1'].border = bordaInferior
    planilha['H1'].border = bordaInferior

    planilha['A1'].fill = preenchimentoAzul
    planilha['B1'].fill = preenchimentoAzul
    planilha['C1'].fill = preenchimentoAzul
    planilha['D1'].fill = preenchimentoAzul
    planilha['E1'].fill = preenchimentoAzul
    planilha['F1'].fill = preenchimentoAzul
    planilha['G1'].fill = preenchimentoAzul
    planilha['H1'].fill = preenchimentoAzul


    planilha.column_dimensions["A"].width = 20.0
    planilha.column_dimensions["B"].width = 20.0
    planilha.column_dimensions["C"].width = 20.0
    planilha.column_dimensions["D"].width = 20.0
    planilha.column_dimensions["E"].width = 16.0
    planilha.column_dimensions["F"].width = 20.0
    planilha.column_dimensions["G"].width = 20.0
    planilha.column_dimensions["H"].width = 8.0

    for item in lista:
        for produto in produtos:
            if item.produto.id == produto.id:
                descricao = produto.descricao.replace(';','')
                descricao = descricao.splitlines()
                print(descricao)
                planilha['A' + str(cont)] = produto.nome
                planilha['B' + str(cont)] = produto.modelo
                planilha['C' + str(cont)] = produto.fabricante
                planilha['D' + str(cont)] = produto.fornecedor.razao_social
                planilha['E' + str(cont)] = item.quantidade
                planilha['F' + str(cont)] = produto.valor
                planilha['G' + str(cont)] = format(produto.data, "%d/%m/%Y")
                planilha['H' + str(cont)] = '= E' + str(cont) + '*F' + str(cont)

                # print(format(produto.data, "%d/%m/%Y"))

                planilha['A' + str(cont)].font = ft_item
                planilha['B' + str(cont)].font = ft_item
                planilha['C' + str(cont)].font = ft_item
                planilha['D' + str(cont)].font = ft_item
                planilha['E' + str(cont)].font = ft_item
                planilha['F' + str(cont)].font = ft_item_italico
                planilha['G' + str(cont)].font = ft_item
                planilha['H' + str(cont)].font = ft_item_negrito

                planilha['A' + str(cont)].alignment = alinhamentoEsquerda
                planilha['B' + str(cont)].alignment = alinhamentoEsquerda
                planilha['C' + str(cont)].alignment = alinhamentoEsquerda
                planilha['D' + str(cont)].alignment = alinhamentoEsquerda
                planilha['E' + str(cont)].alignment = alinhamento
                planilha['F' + str(cont)].alignment = alinhamento
                planilha['G' + str(cont)].alignment = alinhamento
                planilha['H' + str(cont)].alignment = alinhamento

                if cont % 2 == 0:
                    planilha['A' + str(cont)].fill = preenchimentoCinza
                    planilha['B' + str(cont)].fill = preenchimentoCinza
                    planilha['C' + str(cont)].fill = preenchimentoCinza
                    planilha['D' + str(cont)].fill = preenchimentoCinza
                    planilha['E' + str(cont)].fill = preenchimentoCinza
                    planilha['F' + str(cont)].fill = preenchimentoCinza
                    planilha['G' + str(cont)].fill = preenchimentoCinza
                    planilha['H' + str(cont)].fill = preenchimentoCinza
                else:
                    planilha['A' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['B' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['C' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['D' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['E' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['F' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['G' + str(cont)].fill = preenchimentoAzulClaro
                    planilha['H' + str(cont)].fill = preenchimentoAzulClaro

        cont += 1

    planilha['H'+str(cont)] = '= SUM(H2:H' + str(cont-1) +')'
    planilha['H' + str(cont)].font = ft_item_negrito
    planilha['H' + str(cont)].alignment = alinhamento
    planilha['H' + str(cont)].fill = preenchimentoVerde

    planilha['A' + str(cont)].border = bordaSuperior
    planilha['B' + str(cont)].border = bordaSuperior
    planilha['C' + str(cont)].border = bordaSuperior
    planilha['D' + str(cont)].border = bordaSuperior
    planilha['E' + str(cont)].border = bordaSuperior
    planilha['F' + str(cont)].border = bordaSuperior
    planilha['G' + str(cont)].border = bordaSuperior
    planilha['H' + str(cont)].border = bordaSuperior


    wb.save('documents/documents/media/' +nome_doc+'.xlsx')
    gerar_doc(nome_doc)

    return redirect('lista_contatos')

def gerar_docx(nome_doc):
    produtos = Produtos.objects.order_by('nome')
    lista = ListaMaterial.objects.order_by('produto__nome')
    grupos = Grupos.objects.all()
    cgrupo = 1

    doc = docx.Document()
    # run = doc.add_paragraph().add_run()
    # '''Apply style'''
    # style = doc.styles['Normal']
    # font = style.font
    # font.name = 'Calibri'
    # font.size = docx.shared.Pt(11)

    doc.add_paragraph('ANEXO A – ESPECIFICAÇÕES TÉCNICAS').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[0].runs[0].font.size = Pt(16)
    doc.paragraphs[0].runs[0].font.bold = True
    doc.paragraphs[0].runs[0].font.name = 'Calibri'
    doc.add_paragraph('TODOS OS PRODUTOS OFERTADOS DEVERÃO TER, EM SUA COMPOSIÇÃO DE CUSTOS, OS VALORES REFERENTE A INSTALAÇÃO.').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.paragraphs[0].runs[0].font.size = Pt(16)
    doc.paragraphs[1].runs[0].font.size = Pt(11)
    doc.paragraphs[1].runs[0].font.name = 'Calibri'
    doc.add_paragraph('UPI – UNIDADE DE PONTO DE INFRAESTRUTURA').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.paragraphs[2].runs[0].font.size = Pt(13)
    doc.paragraphs[2].runs[0].font.bold = True
    doc.paragraphs[2].runs[0].font.name = 'Calibri'
    doc.paragraphs[2].runs[0].font.color.rgb = RGBColor(79, 129, 189)



    for grupo in grupos:
        if group_check(grupo.nome):
            contador = 1
            index_group = doc.add_paragraph()
            titulo = ('GRUPO ', str(cgrupo),' ', grupo.nome)
            nome = index_group.add_run(titulo)
            nome.font.bold = True
            nome.font.color.rgb = RGBColor(79, 129, 189)
            for item in lista:
                for produto in produtos:
                    if item.produto.id == produto.id and str(produto.grupo) == grupo.nome:
                        print(str(produto.grupo) == '')
                        run = doc.add_paragraph().add_run()

                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Calibri'
                        font.size = docx.shared.Pt(11)
                        descricao = produto.descricao
                        descricao = descricao.splitlines()
                        titulo = (str(cgrupo)+'.' + str(contador) + ' ' + produto.nome)
                        print(titulo)
                        paragrafo = doc.add_paragraph()
                        nome = paragrafo.add_run((titulo))
                        nome.font.bold = True
                        nome.font.color.rgb = RGBColor(79, 129, 189)
                        print(descricao)
                        contador += 1
                        for topico in descricao:
                            doc.add_paragraph(
                                (topico), style='List Bullet'
                            )
            cgrupo += 1


    doc.save('documents/documents/media/' + nome_doc +'.docx')
    gerar_doc(nome_doc)

    return redirect('adicionar_lista')


# FUNÇÃO PARA CHECAR SE HÁ ITENS NO GRUPO SOLICITADO.
def group_check(grupo):
    produtos = Produtos.objects.order_by('nome')
    lista = ListaMaterial.objects.order_by('produto__nome')
    for item in lista:
        for produto in produtos:
            if item.produto.id == produto.id and str(produto.grupo) == grupo:
                return True

    return False

def gerar_doc(nome):
    try:
        f = File(open(os.path.join(settings.MEDIA_ROOT,'documents/media/' + nome +'.docx'),'rb'))
    except:
        f = File(open(os.path.join(settings.MEDIA_ROOT, 'documents/media/' + nome + '.xlsx'), 'rb'))
    doc = DocFiles()
    doc.docupload= f

    doc.title = nome

    doc.save(nome)




def download_doc(path):
    file_path=os.path.join(settings.MEDIA_ROOT,path)
    if os.path.exists(file_path):
        with open(file_path,'rb') as fh:
            response = HttpResponse(fh.read(),content_type="aplication/docupload" )
            response['Content-Disposition'] = 'inline;  filename='+os.path.basename(file_path)
            return response
        raise Http404

def listar_download(request):
    files = {'files' : DocFiles.objects.all()}
    return render(request,'download_list.html', files)

def get_name_docx(request):
    if request.method == 'POST':
        form = NameForm(request.POST)
        if form.is_valid():
            nome = form.cleaned_data['project_name']
            gerar_docx(nome)
            return redirect('listar_downloads')

    else:
        form = NameForm()
    return render(request,'formulario_docs.html', {'form' : form})

def get_name_xlsx(request):
    if request.method == 'POST':
        form = NameForm(request.POST)
        if form.is_valid():
            nome = form.cleaned_data['project_name']
            gerar_xlsx(nome)
            return redirect('listar_downloads')

    else:
        form = NameForm()
    return render(request,'formulario_docs.html', {'form' : form})

def delete_doc(request,id):

    documento = get_object_or_404(DocFiles, id =  id)
    print(str(documento.docupload))
    try:
        os.remove('documents/documents/media/' + documento.title +'.docx')
    except:
        os.remove('documents/documents/media/' + documento.title + '.xlsx')
    os.remove(str(documento.docupload))
    DocFiles.objects.filter(id=id).delete()
    return redirect('listar_downloads')